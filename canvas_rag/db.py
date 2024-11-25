# db.py
import sqlite3
import chromadb
from datetime import datetime
from typing import Dict, List, Any, Optional
import PyPDF2
import io
from docx import Document
from bs4 import BeautifulSoup


class CanvasDatabase:
    def __init__(self):
        # Initialize SQLite
        # sqlite_conn: holds connection object to SQLite database
        # cursor: executes SQL commands and fetches results from the database
        self.sqlite_conn = sqlite3.connect('canvas.db')
        self.cursor = self.sqlite_conn.cursor()
        
        # Initialize ChromaDB for vector search
        # chroma_client: performs operations for vector storage and retrieval
        # collection: stores and manages the vector representations of the Canvas content
        self.chroma_client = chromadb.Client()
        self.collection = self.chroma_client.get_or_create_collection(
            name="canvas_content"
        )
        
        # Create all tables
        self.init_tables()

    def init_tables(self):
        """Initialize comprehensive database structure for all Canvas data"""
        # executescript() executes batches of SQL commands
        self.cursor.executescript('''
        -- Courses table
        CREATE TABLE IF NOT EXISTS courses (
            id TEXT PRIMARY KEY,
            name TEXT NOT NULL,
            code TEXT,
            term_id TEXT,
            start_date TEXT,
            end_date TEXT,
            description TEXT,
            syllabus_content TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );

        -- Assignments table
        CREATE TABLE IF NOT EXISTS assignments (
            id TEXT PRIMARY KEY,
            course_id TEXT,
            title TEXT NOT NULL,
            description TEXT,
            due_date TEXT,
            points_possible REAL,
            submission_types TEXT,
            grading_type TEXT,
            position INTEGER,
            group_category_id TEXT,
            allowed_attempts INTEGER,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (course_id) REFERENCES courses (id)
        );

        -- Modules table
        CREATE TABLE IF NOT EXISTS modules (
            id TEXT PRIMARY KEY,
            course_id TEXT,
            name TEXT NOT NULL,
            position INTEGER,
            prerequisites TEXT,
            unlock_date TEXT,
            require_sequential_progress BOOLEAN,
            published BOOLEAN,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (course_id) REFERENCES courses (id)
        );

        -- Module Items table
        CREATE TABLE IF NOT EXISTS module_items (
            id TEXT PRIMARY KEY,
            module_id TEXT,
            title TEXT NOT NULL,
            type TEXT,
            content_id TEXT,
            position INTEGER,
            indent_level INTEGER,
            external_url TEXT,
            completion_requirement TEXT,
            published BOOLEAN,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (module_id) REFERENCES modules (id)
        );

        -- Announcements table
        CREATE TABLE IF NOT EXISTS announcements (
            id TEXT PRIMARY KEY,
            course_id TEXT,
            title TEXT NOT NULL,
            message TEXT,
            author_id TEXT,
            posted_date TEXT,
            delayed_post_date TEXT,
            published BOOLEAN,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (course_id) REFERENCES courses (id)
        );

        -- Discussion Posts table
        CREATE TABLE IF NOT EXISTS discussion_posts (
            id TEXT PRIMARY KEY,
            discussion_id TEXT,
            user_id TEXT,
            parent_id TEXT,
            message TEXT,
            posted_date TEXT,
            edited_date TEXT,
            rating_count INTEGER,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (discussion_id) REFERENCES discussions (id),
            FOREIGN KEY (parent_id) REFERENCES discussion_posts (id)
        );

        -- Files table
        CREATE TABLE IF NOT EXISTS files (
            id TEXT PRIMARY KEY,
            course_id TEXT,
            filename TEXT NOT NULL,
            display_name TEXT,
            content_type TEXT,
            file_size INTEGER,
            folder_path TEXT,
            file_data BLOB,
            extracted_text TEXT,
            url TEXT,
            thumbnail_url TEXT,
            modified_date TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (course_id) REFERENCES courses (id)
        );

        -- Pages table
        CREATE TABLE IF NOT EXISTS pages (
            id TEXT PRIMARY KEY,
            course_id TEXT,
            title TEXT NOT NULL,
            body TEXT,
            editing_roles TEXT,
            published BOOLEAN,
            front_page BOOLEAN,
            todo_date TEXT,
            edited_date TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (course_id) REFERENCES courses (id)
        );

        -- Submissions table
        CREATE TABLE IF NOT EXISTS submissions (
            id TEXT PRIMARY KEY,
            assignment_id TEXT,
            user_id TEXT,
            submitted_date TEXT,
            grade TEXT,
            score REAL,
            feedback TEXT,
            submission_type TEXT,
            submission_data TEXT,
            late BOOLEAN,
            missing BOOLEAN,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (assignment_id) REFERENCES assignments (id)
        );

        -- Calendar Events table
        CREATE TABLE IF NOT EXISTS calendar_events (
            id TEXT PRIMARY KEY,
            course_id TEXT,
            title TEXT NOT NULL,
            description TEXT,
            start_date TEXT,
            end_date TEXT,
            location TEXT,
            all_day BOOLEAN,
            recurring BOOLEAN,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (course_id) REFERENCES courses (id)
        );

        -- Grades table
        CREATE TABLE IF NOT EXISTS grades (
            id TEXT PRIMARY KEY,
            course_id TEXT,
            assignment_id TEXT,
            user_id TEXT,
            score REAL,
            grade TEXT,
            graded_at TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (course_id) REFERENCES courses (id),
            FOREIGN KEY (assignment_id) REFERENCES assignments (id)
        );

        -- Create necessary indices for better performance
        CREATE INDEX IF NOT EXISTS idx_assignments_course 
        ON assignments (course_id);
        
        CREATE INDEX IF NOT EXISTS idx_modules_course 
        ON modules (course_id);
        
        CREATE INDEX IF NOT EXISTS idx_pages_course 
        ON pages (course_id);
        
        CREATE INDEX IF NOT EXISTS idx_files_course 
        ON files (course_id);
        
        CREATE INDEX IF NOT EXISTS idx_submissions_assignment 
        ON submissions (assignment_id);
        
        CREATE INDEX IF NOT EXISTS idx_grades_course 
        ON grades (course_id);
        
        CREATE INDEX IF NOT EXISTS idx_grades_assignment 
        ON grades (assignment_id);
        
        CREATE INDEX IF NOT EXISTS idx_calendar_events_course 
        ON calendar_events (course_id);
        
        CREATE INDEX IF NOT EXISTS idx_announcements_course 
        ON announcements (course_id);
        ''')
        # saves changes to database
        self.sqlite_conn.commit()

    def store_canvas_item(self, item_type: str, data: Dict[str, Any]) -> bool:
        # takes in a string of the item type and a dictionary containing the item's data
        # returns boolean if the item was stored properly
        """Store any type of Canvas item"""
        try:
            # Check if the item already exists in the database
            if self._item_exists(item_type, data['id']):
                print(f"{item_type} with ID {data['id']} already exists. Skipping insertion.")
                return False

            # Stores data in SQLite database, involves executing an INSERT SQL command
            self._store_in_sqlite(item_type, data)

            # Prepare content for ChromaDB
            content = self._prepare_content(item_type, data)

            # If content is prepared, add it to ChromaDB
            if content:
                self.collection.add(
                    documents=[content],
                    metadatas=[{
                        "type": item_type,
                        "item_id": data['id'],
                        "course_id": data.get('course_id'),
                        "timestamp": datetime.now().isoformat()
                    }],
                    ids=[f"{item_type}_{data['id']}"]
                )

            # Commit changes to the SQLite database
            self.sqlite_conn.commit()
            return True

        except Exception as e:
            print(f"Error storing {item_type}: {e}")
            self.sqlite_conn.rollback()
            return False

    def _item_exists(self, item_type: str, item_id: str) -> bool:
        """
        Check if an item already exists in the database
        Args:
            item_type: Type of canvas item
            item_id: ID of the item
        Returns:
            bool: True if the item exists, False otherwise
        """
        table_name = {
            'course': 'courses',
            'assignment': 'assignments',
            'announcement': 'announcements',
            'discussion': 'discussions',
            'page': 'pages',
            'module': 'modules',
            'file': 'files',
            'submission': 'submissions',
            'event': 'calendar_events',
            'grade': 'grades'
        }.get(item_type)

        if not table_name:
            raise ValueError(f"Unknown item type: {item_type}")

        self.cursor.execute(f"SELECT 1 FROM {table_name} WHERE id = ?", (item_id,))
        return self.cursor.fetchone() is not None

    def store_file(self, course_id: str, file_name: str, file_data: bytes, content_type: str) -> Optional[str]:
        # stores files like PDFs, images, etc...
        # takes in course_id, file_name, content_type as strings and file_data as binary bytes
        # returns the ID of the stored file as a string or None if the operation fails
        """Store file with extracted text"""
        try:
            # generated a unique file ID based off course ID and current timestamp
            file_id = f"file_{course_id}_{datetime.now().strftime('%Y%m%d%H%M%S')}"
            # processes binary data and returns extracted text, which is useful for searching and indexing
            extracted_text = self._extract_file_text(file_data, content_type)

            # executes an SQL INSERT command to add the file's details in the files table of the database
            self.cursor.execute('''
                INSERT INTO files (id, course_id, filename, content_type, file_size, file_data, extracted_text)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (
                file_id,
                course_id,
                file_name,
                content_type,
                len(file_data),
                file_data,
                extracted_text
            ))

            # if there is extracted_text, it is added to the ChromaDB collection
            if extracted_text:
                self.collection.add(
                    documents=[extracted_text],
                    metadatas=[{
                        "type": "file",
                        "file_id": file_id,
                        "course_id": course_id,
                        "filename": file_name,
                        "content_type": content_type,
                        "timestamp": datetime.now().isoformat()
                    }],
                    ids=[file_id]
                )

            self.sqlite_conn.commit()
            return file_id

        except Exception as e:
            print(f"Error storing file: {e}")
            self.sqlite_conn.rollback()
            return None

    def query_content(self, query_text: str, n_results: int = 5) -> List[Dict[str, Any]]:
        # takes in query_text as a string and n_results as an integer
        # returns a list of dictioanaries that contains the data of the queried content
        """Query content using ChromaDB's similarity search"""
        try:
            # performs a similarity search based on provided text query and returns matching documents along with their metadata and IDs
            results = self.collection.query(
                query_texts=[query_text],
                n_results=n_results
            )
            
            formatted_results = []
            for i in range(len(results['documents'][0])):
                metadata = results['metadatas'][0][i]
                
                # Get additional details from SQLite
                details = self._get_item_details(
                    metadata['type'],
                    metadata.get('item_id') or metadata.get('file_id')
                )
                
                formatted_results.append({
                    'content': results['documents'][0][i],
                    'metadata': metadata,
                    'details': details,
                    'id': results['ids'][0][i]
                })
            
            return formatted_results

        except Exception as e:
            print(f"Error querying content: {e}")
            return []

    def _prepare_content(self, item_type: str, data: Dict[str, Any]) -> str:
        """
        Prepare content for ChromaDB based on item type
        Args:
            item_type: Type of canvas item (e.g., 'course', 'assignment')
            data: Dictionary containing item data
        Returns:
            Formatted string content for storage in ChromaDB
        """
        content_parts = []
        
        # Prepare content based on the type of item
        if item_type == 'course':
            content_parts.extend([
                f"Course: {data['name']}",
                data.get('description', ''),
                data.get('syllabus_content', '')
            ])
        
        elif item_type == 'assignment':
            content_parts.extend([
                f"Assignment: {data['title']}",
                data.get('description', ''),
                f"Due Date: {data.get('due_date', '')}"
            ])
        
        elif item_type == 'announcement':
            content_parts.extend([
                f"Announcement: {data['title']}",
                data.get('message', '')
            ])
        
        elif item_type == 'discussion':
            content_parts.extend([
                f"Discussion: {data['title']}",
                data.get('message', '')
            ])
        
        elif item_type == 'page':
            content_parts.extend([
                f"Page: {data['title']}",
                data.get('body', '')
            ])
        
        # Filter out empty strings and join with newlines
        return '\n'.join(part for part in content_parts if part)

    def _store_in_sqlite(self, item_type: str, data: Dict[str, Any]):
        """
        Store item in appropriate SQLite table
        Args:
            item_type: Type of canvas item
            data: Dictionary containing item data
        Raises:
            ValueError: If item_type is unknown
        """
        # Map item type to corresponding table name
        table_name = {
            'course': 'courses',
            'assignment': 'assignments',
            'announcement': 'announcements',
            'discussion': 'discussions',
            'page': 'pages',
            'module': 'modules',
            'file': 'files',
            'submission': 'submissions',
            'event': 'calendar_events',
            'grade': 'grades'
        }.get(item_type)

        if not table_name:
            raise ValueError(f"Unknown item type: {item_type}")

        # Get column names for the table
        self.cursor.execute(f"PRAGMA table_info({table_name})")
        columns = [row[1] for row in self.cursor.fetchall()]
        columns.remove('created_at')  # Remove auto-generated column

        # Filter data to match table columns
        filtered_data = {k: v for k, v in data.items() if k in columns}
        
        # Generate SQL for inserting or replacing data
        placeholders = ','.join(['?' for _ in filtered_data])
        columns_str = ','.join(filtered_data.keys())
        sql = f"INSERT OR REPLACE INTO {table_name} ({columns_str}) VALUES ({placeholders})"
        
        # Execute the SQL command
        self.cursor.execute(sql, list(filtered_data.values()))

    def _get_item_details(self, item_type: str, item_id: str) -> Optional[Dict[str, Any]]:
        """
        Get detailed information about an item from SQLite
        Args:
            item_type: Type of canvas item
            item_id: ID of the item
        Returns:
            Dictionary containing item details or None if not found
        """
        # Map item type to corresponding table name
        table_name = {
            'course': 'courses',
            'assignment': 'assignments',
            'announcement': 'announcements',
            'discussion': 'discussions',
            'page': 'pages'
        }.get(item_type)

        if not table_name:
            return None

        # Query the database for the item details
        self.cursor.execute(f"SELECT * FROM {table_name} WHERE id = ?", (item_id,))
        result = self.cursor.fetchone()
        
        if result:
            # Convert the result to a dictionary
            columns = [description[0] for description in self.cursor.description]
            return dict(zip(columns, result))
        
        return None

    def _extract_file_text(self, file_data: bytes, content_type: str) -> str:
        """
        Extract text content from file data
        Args:
            file_data: Binary file data
            content_type: MIME type of the file (e.g., 'application/pdf', 'text/plain')
        Returns:
            str: Extracted text content or error message
        """
        try:
            # Create file-like object from bytes
            file_obj = io.BytesIO(file_data)
            
            # Handle different file types
            if content_type == 'application/pdf':
                try:
                    # Handle PDF files
                    pdf_reader = PyPDF2.PdfReader(file_obj)
                    text_content = []
                    for page in pdf_reader.pages:
                        text_content.append(page.extract_text())
                    return '\n'.join(text_content)
                except Exception as e:
                    print(f"Error processing PDF: {e}")
                    return "Error: Unable to process PDF file"
                
            elif content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                try:
                    # Handle Word documents
                    doc = Document(file_obj)
                    # Extract text from paragraphs
                    paragraphs = [paragraph.text for paragraph in doc.paragraphs]
                    # Extract text from tables
                    tables = []
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = ' | '.join(cell.text for cell in row.cells)
                            if row_text.strip():  # Only add non-empty rows
                                tables.append(row_text)
                    
                    return '\n'.join(paragraphs + tables)
                except Exception as e:
                    print(f"Error processing Word document: {e}")
                    return "Error: Unable to process Word document"
                
            elif content_type == 'text/plain':
                try:
                    # Handle plain text files
                    return file_data.decode('utf-8', errors='ignore')
                except Exception as e:
                    print(f"Error processing text file: {e}")
                    return "Error: Unable to process text file"
                
            elif content_type == 'text/markdown' or content_type == 'text/md':
                try:
                    # Handle markdown files - treat as plain text
                    return file_data.decode('utf-8', errors='ignore')
                except Exception as e:
                    print(f"Error processing markdown file: {e}")
                    return "Error: Unable to process markdown file"
                
            elif content_type == 'text/html':
                try:
                    # Handle HTML files
                    html_text = file_data.decode('utf-8', errors='ignore')
                    soup = BeautifulSoup(html_text, 'html.parser')
                    # Remove script and style elements
                    for script in soup(["script", "style"]):
                        script.decompose()
                    # Get text
                    text = soup.get_text(separator='\n', strip=True)
                    return text
                except Exception as e:
                    print(f"Error processing HTML file: {e}")
                    return "Error: Unable to process HTML file"
                
            elif content_type.startswith('image/'):
                # For images, store metadata only
                size_kb = len(file_data) / 1024
                return f"Image file ({content_type}) - Size: {size_kb:.2f}KB"
                
            else:
                # For unsupported types, store basic metadata
                size_kb = len(file_data) / 1024
                return f"Unsupported file type: {content_type} - Size: {size_kb:.2f}KB"

        except Exception as e:
            print(f"General error extracting text: {e}")
            return f"Error: Unable to process file ({content_type})"

    def close(self):
        """Close database connections"""
        self.sqlite_conn.close()

if __name__ == "__main__":
    # Test code here
    db = CanvasDatabase()
    print("Database initialized successfully")